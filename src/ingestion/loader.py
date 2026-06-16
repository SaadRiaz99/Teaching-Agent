from pathlib import Path
from typing import List, Optional
from dataclasses import dataclass, field


@dataclass
class Document:
    text: str
    metadata: dict = field(default_factory=dict)
    source: Optional[str] = None


class DocumentLoader:
    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv"}

    def load(self, path: str) -> List[Document]:
        p = Path(path)
        if not p.exists():
            return []
        if p.is_dir():
            return self._load_directory(p)
        return [self._load_file(p)]

    def _load_directory(self, directory: Path) -> List[Document]:
        docs = []
        for ext in self.SUPPORTED_EXTENSIONS:
            for fpath in sorted(directory.rglob(f"*{ext}")):
                try:
                    docs.append(self._load_file(fpath))
                except Exception as e:
                    print(f"  Skipping {fpath}: {e}")
        return docs

    def _load_file(self, path: Path) -> Document:
        suffix = path.suffix.lower()
        metadata = {"source": str(path), "filename": path.name, "extension": suffix}
        if suffix == ".pdf":
            return self._load_pdf(path, metadata)
        elif suffix == ".docx":
            return self._load_docx(path, metadata)
        elif suffix == ".md":
            return self._load_text(path, metadata, is_markdown=True)
        elif suffix in (".txt", ".csv"):
            return self._load_text(path, metadata)
        else:
            raise ValueError(f"Unsupported format: {suffix}")

    def _load_pdf(self, path: Path, metadata: dict) -> Document:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        metadata["page_count"] = len(reader.pages)
        return Document(text=text, metadata=metadata, source=str(path))

    def _load_docx(self, path: Path, metadata: dict) -> Document:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        return Document(text=text, metadata=metadata, source=str(path))

    def _load_text(self, path: Path, metadata: dict, is_markdown: bool = False) -> Document:
        text = path.read_text(encoding="utf-8")
        if is_markdown:
            import markdown
            text = markdown.markdown(text)
            import re
            text = re.sub(r"<[^>]+>", "", text)
        metadata["is_markdown"] = is_markdown
        return Document(text=text, metadata=metadata, source=str(path))
